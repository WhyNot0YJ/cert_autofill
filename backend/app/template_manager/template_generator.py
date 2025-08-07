#!/usr/bin/env python3
"""
Template Generator
模板生成器 - 根据用户配置生成自定义模板
"""

from docx import Document
from docxtpl import DocxTemplate
import os
import shutil
import json
from datetime import datetime
from typing import Dict, List, Any, Optional

class TemplateGenerator:
    """模板生成器类"""
    
    def __init__(self, template_dir: str = "backend/templates"):
        self.template_dir = template_dir
        self.ensure_template_dir()
    
    def ensure_template_dir(self):
        """确保模板目录存在"""
        os.makedirs(self.template_dir, exist_ok=True)
    
    def get_available_variables(self) -> Dict[str, str]:
        """获取所有可用的模板变量"""
        return {
            # 基础信息变量
            'approval_no': '批准号',
            'information_folder_no': '信息文件夹号',
            'company_name': '公司名称',
            'company_address': '公司地址',
            
            # 技术参数变量
            'windscreen_thick': '风窗厚度',
            'interlayer_thick': '夹层厚度',
            'glass_layers': '玻璃层数',
            'interlayer_layers': '夹层数',
            'interlayer_type': '夹层类型',
            'glass_treatment': '玻璃处理',
            'coating_type': '涂层类型',
            'coating_thick': '涂层厚度',
            'coating_color': '涂层颜色',
            'material_nature': '材料性质',
            'safety_class': '安全等级',
            'pane_desc': '玻璃板描述',
            
            # 车辆信息变量
            'veh_cat': '车辆类别',
            'veh_type': '车辆类型',
            'veh_mfr': '车辆制造商',
            'dev_area': '开发区域',
            'seg_height': '段高度',
            'curv_radius': '曲率半径',
            'inst_angle': '安装角度',
            'seat_angle': '座椅角度',
            'rpoint_coords': '参考点坐标',
            'dev_desc': '开发描述',
            'remarks': '备注',
            
            # 图片变量
            'company_logo': '公司logo',
            'certification_logo': '认证logo',
            'page_image': '页面图片',
            
            # 页面信息
            'page_info': '页面信息'
        }
    
    def create_template_config(self, template_name: str, selected_variables: List[str], 
                             template_description: str = "") -> Dict[str, Any]:
        """创建模板配置"""
        config = {
            "template_name": template_name,
            "description": template_description,
            "created_at": datetime.now().isoformat(),
            "variables": selected_variables,
            "all_variables": self.get_available_variables(),
            "template_file": f"{template_name}.docx",
            "config_file": f"{template_name}_config.json"
        }
        return config
    
    def generate_template(self, template_name: str, selected_variables: List[str], 
                         template_description: str = "", source_template: str = None) -> Dict[str, Any]:
        """生成自定义模板"""
        try:
            # 创建模板配置
            config = self.create_template_config(template_name, selected_variables, template_description)
            
            # 确定源模板文件
            if source_template and os.path.exists(source_template):
                source_file = source_template
            else:
                # 使用默认模板
                source_file = "backend/client/IF_Template_2.docx"
                if not os.path.exists(source_file):
                    # 创建一个基础模板
                    source_file = self.create_base_template()
            
            # 生成目标模板文件路径
            target_file = os.path.join(self.template_dir, f"{template_name}.docx")
            
            # 复制源模板
            shutil.copy2(source_file, target_file)
            
            # 保存配置
            config_file = os.path.join(self.template_dir, f"{template_name}_config.json")
            with open(config_file, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
            
            return {
                "success": True,
                "message": f"模板 '{template_name}' 生成成功",
                "data": {
                    "template_name": template_name,
                    "template_file": target_file,
                    "config_file": config_file,
                    "variables": selected_variables,
                    "description": template_description
                }
            }
            
        except Exception as e:
            return {
                "success": False,
                "message": f"模板生成失败: {str(e)}",
                "error": str(e)
            }
    
    def create_base_template(self) -> str:
        """创建基础模板文件"""
        doc = Document()
        doc.add_heading('证书模板', 0)
        doc.add_paragraph('这是一个基础模板，请根据需要进行自定义。')
        
        # 添加一些基础变量示例
        doc.add_heading('基础信息', level=1)
        doc.add_paragraph('批准号: {{approval_no}}')
        doc.add_paragraph('信息文件夹号: {{information_folder_no}}')
        doc.add_paragraph('公司名称: {{company_name}}')
        
        doc.add_heading('技术参数', level=1)
        doc.add_paragraph('风窗厚度: {{windscreen_thick}}')
        doc.add_paragraph('夹层厚度: {{interlayer_thick}}')
        doc.add_paragraph('玻璃层数: {{glass_layers}}')
        
        # 保存基础模板
        base_template_path = os.path.join(self.template_dir, "base_template.docx")
        doc.save(base_template_path)
        return base_template_path
    
    def list_templates(self) -> List[Dict[str, Any]]:
        """列出所有可用模板"""
        templates = []
        
        if not os.path.exists(self.template_dir):
            return templates
        
        for file in os.listdir(self.template_dir):
            if file.endswith('.docx') and not file.startswith('~$'):
                template_name = file.replace('.docx', '')
                config_file = os.path.join(self.template_dir, f"{template_name}_config.json")
                
                template_info = {
                    "name": template_name,
                    "file": file,
                    "path": os.path.join(self.template_dir, file),
                    "created_at": datetime.fromtimestamp(os.path.getctime(os.path.join(self.template_dir, file))).isoformat(),
                    "size": os.path.getsize(os.path.join(self.template_dir, file))
                }
                
                # 读取配置信息
                if os.path.exists(config_file):
                    try:
                        with open(config_file, 'r', encoding='utf-8') as f:
                            config = json.load(f)
                            template_info.update(config)
                    except:
                        pass
                
                templates.append(template_info)
        
        return templates
    
    def delete_template(self, template_name: str) -> Dict[str, Any]:
        """删除模板"""
        try:
            template_file = os.path.join(self.template_dir, f"{template_name}.docx")
            config_file = os.path.join(self.template_dir, f"{template_name}_config.json")
            
            deleted_files = []
            
            if os.path.exists(template_file):
                os.remove(template_file)
                deleted_files.append(template_file)
            
            if os.path.exists(config_file):
                os.remove(config_file)
                deleted_files.append(config_file)
            
            return {
                "success": True,
                "message": f"模板 '{template_name}' 删除成功",
                "deleted_files": deleted_files
            }
            
        except Exception as e:
            return {
                "success": False,
                "message": f"删除模板失败: {str(e)}",
                "error": str(e)
            }
    
    def get_template_config(self, template_name: str) -> Dict[str, Any]:
        """获取模板配置"""
        config_file = os.path.join(self.template_dir, f"{template_name}_config.json")
        
        if os.path.exists(config_file):
            try:
                with open(config_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                return {"error": f"读取配置失败: {str(e)}"}
        
        return {"error": "配置文件不存在"}

# 创建全局实例
template_generator = TemplateGenerator() 