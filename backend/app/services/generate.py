import os
import json
import shutil
from datetime import datetime
from docxtpl import DocxTemplate
import subprocess
import platform

def generate_document(template_id, fields, output_path):
    """
    根据模板和数据生成文档
    
    Args:
        template_id: 模板ID
        fields: 字段数据
        output_path: 输出文件路径
    
    Returns:
        bool: 生成是否成功
    """
    try:
        # 这里应该从数据库获取模板信息
        # 暂时使用硬编码的模板路径
        template_path = os.path.join(
            os.path.dirname(os.path.dirname(__file__)), 
            '..', 
            'templates', 
            'certificate_template.docx'
        )
        
        # 确保输出目录存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # 如果模板文件不存在，创建一个简单的文本文件
        if not os.path.exists(template_path):
            # 创建简单的文本模板
            template_content = """TÜV NORD 认证证书

证书编号: {certificate_number}
企业名称: {enterprise_name}
企业英文名: {enterprise_english_name}
注册号: {registration_number}
法定代表人: {legal_representative}

认证类型: {certification_type}
产品名称: {product_name}
产品型号: {product_model}
认证范围: {certification_scope}

发证日期: {issue_date}
有效期至: {expiry_date}
发证机构: {issuing_authority}

联系人: {contact_person}
联系电话: {contact_phone}
联系邮箱: {contact_email}
地址: {address}

技术规格参数:
{technical_specifications}

测试标准: {test_standards}
测试结果: {test_results}

备注: {remarks}

本证书证明上述企业已通过TÜV NORD的认证审核，
符合相关标准要求。

TÜV NORD
{generated_date}"""
            
            # 准备上下文数据
            context = {
                **fields,
                'generated_date': datetime.now().strftime('%Y-%m-%d'),
                'generated_time': datetime.now().strftime('%H:%M:%S')
            }
            
            # 格式化内容
            content = template_content.format(**context)
            
            # 保存为文本文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return True
        else:
            # 使用docxtpl处理Word模板
            from docxtpl import DocxTemplate
            doc = DocxTemplate(template_path)
            
            # 准备上下文数据
            context = {
                **fields,
                'generated_date': datetime.now().strftime('%Y-%m-%d'),
                'generated_time': datetime.now().strftime('%H:%M:%S')
            }
            
            # 渲染模板
            doc.render(context)
            
            # 保存文档
            doc.save(output_path)
            return True
            
    except Exception as e:
        print(f"生成文档失败: {e}")
        return False

def generate_pdf_from_docx(docx_path, pdf_path):
    """
    将Word文档转换为PDF
    
    Args:
        docx_path: Word文档路径
        pdf_path: PDF输出路径
    
    Returns:
        bool: 转换是否成功
    """
    try:
        # 确保输出目录存在
        os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
        
        # 检查系统平台
        system = platform.system()
        
        if system == "Windows":
            # Windows系统使用LibreOffice
            libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
            if not os.path.exists(libreoffice_path):
                libreoffice_path = r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
            
            if os.path.exists(libreoffice_path):
                cmd = [
                    libreoffice_path,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", os.path.dirname(pdf_path),
                    docx_path
                ]
                result = subprocess.run(cmd, capture_output=True, text=True)
                
                if result.returncode == 0:
                    # LibreOffice会生成一个PDF文件，但文件名可能不同
                    # 我们需要找到生成的PDF文件并重命名为目标文件名
                    docx_basename = os.path.splitext(os.path.basename(docx_path))[0]
                    generated_pdf = os.path.join(os.path.dirname(pdf_path), f"{docx_basename}.pdf")
                    
                    if os.path.exists(generated_pdf):
                        # 如果生成的文件名与目标文件名不同，则重命名
                        if generated_pdf != pdf_path:
                            shutil.move(generated_pdf, pdf_path)
                        return True
                    else:
                        print(f"PDF文件未生成: {generated_pdf}")
                        return False
                else:
                    print(f"LibreOffice转换失败: {result.stderr}")
                    return False
            else:
                print("LibreOffice not found. Please install LibreOffice for PDF conversion.")
                return False
                
        elif system == "Darwin":  # macOS
            # macOS系统使用LibreOffice
            libreoffice_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
            if os.path.exists(libreoffice_path):
                cmd = [
                    libreoffice_path,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", os.path.dirname(pdf_path),
                    docx_path
                ]
                result = subprocess.run(cmd, capture_output=True, text=True)
                
                if result.returncode == 0:
                    # LibreOffice会生成一个PDF文件，但文件名可能不同
                    # 我们需要找到生成的PDF文件并重命名为目标文件名
                    docx_basename = os.path.splitext(os.path.basename(docx_path))[0]
                    generated_pdf = os.path.join(os.path.dirname(pdf_path), f"{docx_basename}.pdf")
                    
                    if os.path.exists(generated_pdf):
                        # 如果生成的文件名与目标文件名不同，则重命名
                        if generated_pdf != pdf_path:
                            shutil.move(generated_pdf, pdf_path)
                        return True
                    else:
                        print(f"PDF文件未生成: {generated_pdf}")
                        return False
                else:
                    print(f"LibreOffice转换失败: {result.stderr}")
                    return False
            else:
                print("LibreOffice not found. Please install LibreOffice for PDF conversion.")
                return False
                
        elif system == "Linux":
            # Linux系统使用LibreOffice
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", os.path.dirname(pdf_path),
                docx_path
            ]
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode == 0:
                # LibreOffice会生成一个PDF文件，但文件名可能不同
                # 我们需要找到生成的PDF文件并重命名为目标文件名
                docx_basename = os.path.splitext(os.path.basename(docx_path))[0]
                generated_pdf = os.path.join(os.path.dirname(pdf_path), f"{docx_basename}.pdf")
                
                if os.path.exists(generated_pdf):
                    # 如果生成的文件名与目标文件名不同，则重命名
                    if generated_pdf != pdf_path:
                        shutil.move(generated_pdf, pdf_path)
                    return True
                else:
                    print(f"PDF文件未生成: {generated_pdf}")
                    return False
            else:
                print(f"LibreOffice转换失败: {result.stderr}")
                return False
            
        else:
            print(f"Unsupported operating system: {system}")
            return False
            
    except Exception as e:
        print(f"PDF转换失败: {e}")
        return False



def generate_multi_template_document(base_template_path, vehicle_template_path, fields, output_path):
    """
    使用 docxcompose 生成多模板合并文档（推荐方案）
    
    Args:
        base_template_path: 基础模板路径（固定部分）
        vehicle_template_path: 车辆模板路径（循环部分）
        fields: 表单字段数据
        output_path: 输出文件路径
    
    Returns:
        bool: 生成是否成功
    """
    try:
        from docxcompose.composer import Composer
        from docx import Document
        from docxtpl import DocxTemplate
        from docx.enum.text import WD_BREAK
        import tempfile
        import os
        
        print(f"开始使用 docxcompose 生成多模板文档")
        print(f"基础模板: {base_template_path}")
        print(f"车辆模板: {vehicle_template_path}")
        
        # 检查模板文件是否存在
        if not os.path.exists(base_template_path):
            print(f"❌ 基础模板文件不存在: {base_template_path}")
            return False
            
        if not os.path.exists(vehicle_template_path):
            print(f"❌ 车辆模板文件不存在: {vehicle_template_path}")
            return False
        
        # 获取车辆信息
        vehicles = fields.get('vehicles', [])
        print(f"发现 {len(vehicles)} 个车辆信息")
        
        # 准备基础上下文数据
        base_context = {
            **fields,
            'generated_date': datetime.now().strftime('%Y-%m-%d'),
            'generated_time': datetime.now().strftime('%H:%M:%S'),
            'has_vehicles': len(vehicles) > 0,
            'total_vehicles': len(vehicles)
        }
        
        # 如果没有车辆信息，清空车辆相关字段
        if not vehicles:
            base_context.update({
                'veh_mfr': '',
                'veh_type': '',
                'veh_cat': '',
                'dev_area': '',
                'seg_height': '',
                'curv_radius': '',
                'inst_angle': '',
                'seat_angle': '',
                'rpoint_coords': '',
                'dev_desc': ''
            })
        else:
            # 如果有车辆信息，使用第一个车辆的数据填充基础模板变量
            first_vehicle = vehicles[0]
            base_context.update({
                'veh_mfr': first_vehicle.get('veh_mfr', ''),
                'veh_type': first_vehicle.get('veh_type', ''),
                'veh_cat': first_vehicle.get('veh_cat', ''),
                'dev_area': first_vehicle.get('dev_area', ''),
                'seg_height': first_vehicle.get('seg_height', ''),
                'curv_radius': first_vehicle.get('curv_radius', ''),
                'inst_angle': first_vehicle.get('inst_angle', ''),
                'seat_angle': first_vehicle.get('seat_angle', ''),
                'rpoint_coords': first_vehicle.get('rpoint_coords', ''),
                'dev_desc': first_vehicle.get('dev_desc', '')
            })
        
        # 1. 渲染基础模板
        print("1. 渲染基础模板...")
        base_doc = DocxTemplate(base_template_path)
        base_doc.render(base_context)
        
        # 保存渲染后的基础文档到临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_base:
            base_doc.save(temp_base.name)
            temp_base_path = temp_base.name
        
        # 2. 创建主文档（从基础模板开始）
        master_doc = Document(temp_base_path)
        composer = Composer(master_doc)
        
        # 3. 如果有车辆信息，循环添加车辆模板
        if vehicles:
            print(f"2. 开始添加 {len(vehicles)} 个车辆页面...")
            
            for i, vehicle in enumerate(vehicles):
                print(f"   处理车辆 {i+1}/{len(vehicles)}: {vehicle.get('veh_mfr', 'Unknown')}")
                
                # 准备车辆上下文数据
                vehicle_context = {
                    **base_context,
                    'veh_mfr': vehicle.get('veh_mfr', ''),
                    'veh_type': vehicle.get('veh_type', ''),
                    'veh_cat': vehicle.get('veh_cat', ''),
                    'dev_area': vehicle.get('dev_area', ''),
                    'seg_height': vehicle.get('seg_height', ''),
                    'curv_radius': vehicle.get('curv_radius', ''),
                    'inst_angle': vehicle.get('inst_angle', ''),
                    'seat_angle': vehicle.get('seat_angle', ''),
                    'rpoint_coords': vehicle.get('rpoint_coords', ''),
                    'dev_desc': vehicle.get('dev_desc', ''),
                    'vehicle_index': i + 1,
                    'total_vehicles': len(vehicles)
                }
                
                # 渲染车辆模板
                vehicle_doc = DocxTemplate(vehicle_template_path)
                vehicle_doc.render(vehicle_context)
                
                # 保存渲染后的车辆文档到临时文件
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_vehicle:
                    vehicle_doc.save(temp_vehicle.name)
                    temp_vehicle_path = temp_vehicle.name
                
                # 添加到主文档
                sub_doc = Document(temp_vehicle_path)
                
                # 强制每个车辆信息从新页开始（除了第一个）
                if i > 0:
                    # 在车辆文档开头添加分页符
                    if sub_doc.paragraphs:
                        # 在第一个段落前插入分页符
                        first_paragraph = sub_doc.paragraphs[0]
                        first_run = first_paragraph.runs[0] if first_paragraph.runs else first_paragraph.add_run()
                        first_run.add_break(WD_BREAK.PAGE)
                    else:
                        # 如果文档没有段落，创建一个带分页符的段落
                        new_paragraph = sub_doc.add_paragraph()
                        new_paragraph.add_run().add_break(WD_BREAK.PAGE)
                
                composer.append(sub_doc)
                
                # 清理临时文件
                os.unlink(temp_vehicle_path)
        
        # 4. 保存最终文档
        print("3. 保存最终文档...")
        composer.save(output_path)
        
        # 清理临时文件
        os.unlink(temp_base_path)
        
        print(f"✅ 多模板文档生成成功: {output_path}")
        print(f"📄 文档包含: 1个基础页面 + {len(vehicles)}个车辆页面")
        return True
        
    except Exception as e:
        print(f"❌ 多模板文档生成失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def generate_if_template_document(fields, output_path, template_type="IF_Template_Auto"):
    """
    IF模板文档生成 - 使用 docxcompose 多模板合并方案
    
    Args:
        fields: 表单字段数据
        output_path: 输出文件路径
        template_type: 模板类型
    
    Returns:
        bool: 生成是否成功
    """
    try:
        # 获取当前文件所在目录
        current_dir = os.path.dirname(os.path.abspath(__file__))
        backend_dir = os.path.dirname(os.path.dirname(current_dir))
        
        # 模板路径
        base_template_path = os.path.join(backend_dir, 'templates', 'IF_Template_Base.docx')
        vehicle_template_path = os.path.join(backend_dir, 'templates', 'IF_Template_Vehicle.docx')
        
        # 检查模板文件是否存在
        if not os.path.exists(base_template_path) or not os.path.exists(vehicle_template_path):
            print(f"❌ 模板文件不存在:")
            print(f"   基础模板: {base_template_path}")
            print(f"   车辆模板: {vehicle_template_path}")
            return False
        
        # 使用多模板合并方案
        return generate_multi_template_document(base_template_path, vehicle_template_path, fields, output_path)
        
    except Exception as e:
        print(f"❌ IF模板文档生成失败: {e}")
        import traceback
        traceback.print_exc()
        return False
